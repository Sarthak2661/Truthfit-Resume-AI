from io import BytesIO

from docx import Document

from source.loaders import jd_loader, resume_loader
from tests.conftest import FakePdfReader, make_upload


def make_docx_upload(name="resume.docx"):
    document = Document()
    document.add_paragraph("Python data engineer")
    document.add_paragraph("")
    document.add_paragraph("Kafka and PostgreSQL projects")

    buffer = BytesIO()
    document.save(buffer)
    return make_upload(buffer.getvalue(), name)


def test_resume_loader_extracts_txt():
    uploaded = make_upload(b"Python\nSQL\nKafka", "resume.txt")

    assert resume_loader.extract_text_from_file(uploaded) == "Python\nSQL\nKafka"


def test_resume_loader_extracts_docx():
    uploaded = make_docx_upload()

    text = resume_loader.extract_text_from_file(uploaded)

    assert "Python data engineer" in text
    assert "Kafka and PostgreSQL projects" in text


def test_resume_loader_extracts_pdf_with_reader(monkeypatch):
    monkeypatch.setattr(resume_loader, "PdfReader", FakePdfReader)
    uploaded = make_upload(b"%PDF fake", "resume.pdf")

    assert resume_loader.extract_text_from_file(uploaded) == "First PDF page\nSecond PDF page"


def test_resume_loader_rejects_unsupported_file_type():
    uploaded = make_upload(b"data", "resume.csv")

    try:
        resume_loader.extract_text_from_file(uploaded)
    except ValueError as exc:
        assert "Unsupported resume file type" in str(exc)
    else:
        raise AssertionError("Expected unsupported resume file type to raise ValueError")


def test_jd_loader_combines_uploaded_txt_and_pasted_text():
    uploaded = make_upload(b"Uploaded JD text", "job.txt")

    assert jd_loader.extract_jd_text(uploaded, "Pasted JD text") == "Uploaded JD text\n\nPasted JD text"


def test_jd_loader_extracts_docx():
    uploaded = make_docx_upload("job.docx")

    text = jd_loader.extract_jd_text(uploaded, "")

    assert "Python data engineer" in text


def test_jd_loader_extracts_pdf_with_reader(monkeypatch):
    monkeypatch.setattr(jd_loader, "PdfReader", FakePdfReader)
    uploaded = make_upload(b"%PDF fake", "job.pdf")

    assert "Second PDF page" in jd_loader.extract_jd_text(uploaded, "")
